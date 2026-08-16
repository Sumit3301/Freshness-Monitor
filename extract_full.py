import docx
from docx import Document
from lxml import etree

doc = Document(r'C:\Users\ACER\Downloads\Manuscript Food chemistry.docx')

with open(r'd:\POC project\manuscript_full.txt', 'w', encoding='utf-8') as f:
    para_num = 0
    for para in doc.paragraphs:
        para_num += 1
        text = para.text
        style = para.style.name if para.style else "None"
        
        # Get formatting details for each run
        runs_info = []
        for run in para.runs:
            info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
            }
            if run.font.color and run.font.color.rgb:
                info['color'] = str(run.font.color.rgb)
            if run.font.size:
                info['size'] = run.font.size.pt
            if run.font.name:
                info['font'] = run.font.name
            runs_info.append(info)
        
        if text.strip():
            f.write(f"=== PARA {para_num} [Style: {style}] ===\n")
            f.write(f"{text}\n\n")
    
    # Extract tables
    f.write("\n\n=== TABLES ===\n\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"--- TABLE {t_idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns ---\n")
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            f.write(f"  Row {r_idx}: {' | '.join(row_data)}\n")
        f.write("\n")

print("Done writing full manuscript")

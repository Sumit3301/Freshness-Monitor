"""Verify the revised manuscript changes."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

OUTPUT_PATH = r'C:\Users\ACER\Downloads\Manuscript Food chemistry - REVISED.docx'
verify_doc = Document(OUTPUT_PATH)

checks = [
    ("1. Hypothesis in abstract", "we hypothesized"),
    ("2. Hydrophobicity context", "borderline hydrophilic-to-hydrophobic"),
    ("3. Synergistic removed", "combined integration"),
    ("4a. ASTM soil details", "pH 6.8"),
    ("4b. Biodeg fix", "preliminary screening period"),
    ("5. Ammonia details", "sealed 250 mL glass desiccator"),
    ("6. Film adhesion method", "double-sided adhesive tape"),
    ("7. Sheet to smooth", "smooth and coherent"),
    ("8. XPS functional", "radical scavenging"),
    ("9. Zone of inhibition", "zone of inhibition diameters"),
    ("10. Temperature discussion", "approximately double for every"),
    ("11. Antioxidant diminishing", "diminishing return"),
    ("12. Conclusions fix", "including pH-responsive colorimetric sensing"),
    ("13. Statistical Analysis", "one-way analysis of variance"),
    ("14. NH3 statement", "colour variation upon exposure"),
    ("15. Table 1 footnote", "superscript letters"),
    ("16. Figure 2 magnification", "magnification"),
    ("17. References", "References"),
    ("18. E. coli Gram-negative", "Gram-negative bacterium"),
    ("19. Curcumin description", "enol form"),
    ("21. Headspace clarification", "headspace sensing"),
    ("22. Table 2 discussion", "greatest dynamic colour range"),
    ("23. Novelty justification", "protein-based platform"),
    ("24. Sensing limitations", "limit of detection"),
    ("25. Mechanical limitations", "tensile strength"),
]

print("=" * 70)
print("VERIFICATION RESULTS")
print("=" * 70)

pass_count = 0
fail_count = 0
for label, keyword in checks:
    found = False
    for para in verify_doc.paragraphs:
        if keyword in para.text:
            found = True
            break
    if found:
        print(f"  [PASS] {label}")
        pass_count += 1
    else:
        print(f"  [FAIL] {label} -- keyword '{keyword}' NOT found")
        fail_count += 1

print(f"\nResults: {pass_count} passed, {fail_count} failed out of {len(checks)} checks")

# Also check that "synergistic" is fully removed
syn_count = 0
for para in verify_doc.paragraphs:
    if "synergistic" in para.text.lower():
        syn_count += 1
        print(f"\n  [WARN] 'synergistic' still found in: {para.text[:100]}...")

if syn_count == 0:
    print("\n  [PASS] No remaining 'synergistic' instances found")
else:
    print(f"\n  [WARN] Found {syn_count} remaining 'synergistic' instances")

# Check E. coli incubation temp
for para in verify_doc.paragraphs:
    if "Gram-negative" in para.text and "37" in para.text:
        print(f"\n  [PASS] E. coli incubation at 37C confirmed")
        break

print(f"\nRevised file saved at: {OUTPUT_PATH}")

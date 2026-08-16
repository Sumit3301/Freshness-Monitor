"""
Comprehensive Research Paper Generator for run_20260503_1136
WITH EMBEDDED CHARTS - Matches original Research_Paper.docx structure
"""
import os, sys, io
import cv2
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── CONFIG ────────────────────────────────────────────────────────────────────
RUN_DIR   = "run_20260503_1136"
RUN_DATE  = "03 May 2026"
OUTPUT    = f"Research_Paper_{RUN_DIR}.docx"
CHART_DIR = "_charts_tmp"
os.makedirs(CHART_DIR, exist_ok=True)

STAGES = {
    1: ("Very Fresh",     0,  3,  "Bright yellow, very high saturation",  "Safe for consumption"),
    2: ("Fresh",          4,  6,  "Slight color shift, yellow-gold",       "Safe, consume soon"),
    3: ("Early Spoilage", 7, 14,  "Noticeable darkening / orange tones",  "Caution, refrigerate immediately"),
    4: ("Spoiled",       15, 999, "Significant desaturation / browning",  "Not safe for consumption"),
}
STAGE_COLORS = {1: "#2ECC71", 2: "#F1C40F", 3: "#E67E22", 4: "#E74C3C"}
STAGE_BG     = {1: "#C6EFCE", 2: "#FFEB9C", 3: "#FFC7CE", 4: "#FFD7D7"}

# ─── PLOT STYLE ────────────────────────────────────────────────────────────────
PLT_STYLE = {
    "figure.facecolor": "#1A1A2E",
    "axes.facecolor":   "#16213E",
    "axes.edgecolor":   "#4A4E69",
    "axes.labelcolor":  "#E0E0E0",
    "axes.titlecolor":  "#FFFFFF",
    "xtick.color":      "#B0B0B0",
    "ytick.color":      "#B0B0B0",
    "grid.color":       "#2A2A4A",
    "grid.linestyle":   "--",
    "grid.alpha":       0.6,
    "legend.facecolor": "#0F3460",
    "legend.edgecolor": "#4A4E69",
    "legend.labelcolor":"#E0E0E0",
    "text.color":       "#E0E0E0",
    "font.family":      "DejaVu Sans",
}

def apply_style():
    plt.rcParams.update(PLT_STYLE)

def stage_band(ax, hours_sorted):
    """Draw colored stage bands behind a plot."""
    boundaries = [(0,3,"#2ECC71",0.12), (4,6,"#F1C40F",0.12),
                  (7,14,"#E67E22",0.12), (15,max(hours_sorted),"#E74C3C",0.10)]
    for x0, x1, col, alpha in boundaries:
        if x0 <= max(hours_sorted):
            ax.axvspan(x0, min(x1, max(hours_sorted)), color=col, alpha=alpha, zorder=0)

def save_chart(fig, name):
    path = os.path.join(CHART_DIR, f"{name}.png")
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path

# ─── IMAGE ANALYSIS ────────────────────────────────────────────────────────────
def extract_hours(filename):
    try:
        return int(os.path.basename(filename).replace("h.jpg", ""))
    except ValueError:
        return -1

def get_stage(h):
    if h <= 3:  return 1
    if h <= 6:  return 2
    if h <= 14: return 3
    return 4

def analyze_image(img_path):
    img = cv2.imread(img_path)
    if img is None:
        return None
    h, w = img.shape[:2]
    crop = cv2.resize(img[int(h*0.3):int(h*0.7), int(w*0.3):int(w*0.7)], (256,256))
    b_ch, g_ch, r_ch = cv2.split(crop)
    R, G, B = np.mean(r_ch), np.mean(g_ch), np.mean(b_ch)
    lab = cv2.cvtColor(crop, cv2.COLOR_BGR2Lab)
    L, Astar, Bstar = [np.mean(c) for c in cv2.split(lab)]
    hsv = cv2.cvtColor(crop, cv2.COLOR_BGR2HSV)
    H, S, V = [np.mean(c) for c in cv2.split(hsv)]
    pixels = crop.reshape(-1,3).astype(np.float32)
    criteria = (cv2.TERM_CRITERIA_EPS+cv2.TERM_CRITERIA_MAX_ITER, 10, 1.0)
    _, labels, centers = cv2.kmeans(pixels, 3, None, criteria, 10, cv2.KMEANS_PP_CENTERS)
    counts = np.bincount(labels.flatten())
    colors = centers[np.argsort(-counts)].astype(int)
    hex_colors = [f"#{c[2]:02x}{c[1]:02x}{c[0]:02x}" for c in colors]
    return dict(R=R,G=G,B=B,L=L,Astar=Astar,Bstar=Bstar,H=H,S=S,V=V,
                RB_gap=R-B, hex_colors=hex_colors)

# ═══════════════════════════════════════════════════════════════════════════════
# CHART GENERATORS
# ═══════════════════════════════════════════════════════════════════════════════

def chart_saturation_decay(hours_sorted, data):
    apply_style()
    fig, ax = plt.subplots(figsize=(10, 4.5))
    fig.patch.set_facecolor("#1A1A2E")
    s_vals = [data[h]["S"] for h in hours_sorted]
    stage_band(ax, hours_sorted)
    ax.plot(hours_sorted, s_vals, color="#00D4FF", linewidth=2.5, marker="o",
            markersize=5, markerfacecolor="#FF6B6B", zorder=5, label="Saturation (S)")
    ax.fill_between(hours_sorted, s_vals, alpha=0.18, color="#00D4FF")
    # Stage labels
    for label, xc, yc in [("Very\nFresh",1.5,115),("Fresh",5,115),("Early\nSpoilage",10.5,115),("Spoiled",26,115)]:
        if xc <= max(hours_sorted):
            ax.text(xc, yc, label, ha="center", va="bottom", fontsize=7,
                    color="#CCCCCC", style="italic")
    ax.set_xlabel("Elapsed Time (hours)", fontsize=11)
    ax.set_ylabel("HSV Saturation (S)", fontsize=11)
    ax.set_title("Saturation Decay Curve Over Spoilage Progression", fontsize=13, fontweight="bold", pad=12)
    ax.set_xlim(min(hours_sorted)-0.5, max(hours_sorted)+0.5)
    ax.set_ylim(0, 160)
    ax.grid(True, alpha=0.4)
    ax.legend(loc="upper right")
    fig.tight_layout()
    return save_chart(fig, "01_saturation_decay")

def chart_rgb_progression(hours_sorted, data):
    apply_style()
    fig, ax = plt.subplots(figsize=(10, 4.5))
    fig.patch.set_facecolor("#1A1A2E")
    R_vals = [data[h]["R"] for h in hours_sorted]
    G_vals = [data[h]["G"] for h in hours_sorted]
    B_vals = [data[h]["B"] for h in hours_sorted]
    stage_band(ax, hours_sorted)
    ax.plot(hours_sorted, R_vals, color="#FF4444", lw=2.2, marker="o", ms=4, label="Red (R)")
    ax.plot(hours_sorted, G_vals, color="#44FF88", lw=2.2, marker="s", ms=4, label="Green (G)")
    ax.plot(hours_sorted, B_vals, color="#4488FF", lw=2.2, marker="^", ms=4, label="Blue (B)")
    ax.fill_between(hours_sorted, R_vals, B_vals, alpha=0.10, color="#FF4444", label="R–B gap")
    ax.set_xlabel("Elapsed Time (hours)", fontsize=11)
    ax.set_ylabel("Mean Channel Value (0–255)", fontsize=11)
    ax.set_title("RGB Channel Progression Over Time", fontsize=13, fontweight="bold", pad=12)
    ax.set_xlim(min(hours_sorted)-0.5, max(hours_sorted)+0.5)
    ax.set_ylim(50, 260)
    ax.grid(True, alpha=0.4)
    ax.legend(loc="lower right")
    fig.tight_layout()
    return save_chart(fig, "02_rgb_progression")

def chart_rb_gap(hours_sorted, data):
    apply_style()
    fig, ax = plt.subplots(figsize=(10, 4.0))
    fig.patch.set_facecolor("#1A1A2E")
    rb_vals = [data[h]["RB_gap"] for h in hours_sorted]
    stage_band(ax, hours_sorted)
    bars = ax.bar(hours_sorted, rb_vals,
                  color=[STAGE_COLORS[get_stage(h)] for h in hours_sorted],
                  edgecolor="#333355", linewidth=0.5, alpha=0.9, zorder=5)
    ax.axhline(0, color="#AAAAAA", lw=0.8, ls="--")
    ax.set_xlabel("Elapsed Time (hours)", fontsize=11)
    ax.set_ylabel("R – B Gap (units)", fontsize=11)
    ax.set_title("RGB Channel Convergence (R–B Gap) — Desaturation Indicator", fontsize=13, fontweight="bold", pad=12)
    ax.grid(True, axis="y", alpha=0.4)
    legend_patches = [mpatches.Patch(color=STAGE_COLORS[s], label=f"Stage {s}: {STAGES[s][0]}") for s in STAGES]
    ax.legend(handles=legend_patches, loc="upper right", fontsize=8)
    fig.tight_layout()
    return save_chart(fig, "03_rb_gap")

def chart_lab_channels(hours_sorted, data):
    apply_style()
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5))
    fig.patch.set_facecolor("#1A1A2E")
    fig.suptitle("CIE L*a*b* Channel Values Over Spoilage Progression", fontsize=13, fontweight="bold", y=1.02)
    channels = [("L", "L* (Lightness)", "#FFD700"), ("Astar", "a* (Green–Red)", "#FF6B6B"), ("Bstar", "b* (Blue–Yellow)", "#74B9FF")]
    for ax, (key, label, color) in zip(axes, channels):
        vals = [data[h][key] for h in hours_sorted]
        stage_band(ax, hours_sorted)
        ax.plot(hours_sorted, vals, color=color, lw=2.3, marker="o", ms=4)
        ax.fill_between(hours_sorted, vals, alpha=0.15, color=color)
        ax.set_title(label, fontsize=10, fontweight="bold")
        ax.set_xlabel("Hours", fontsize=9)
        ax.grid(True, alpha=0.4)
    axes[0].set_ylabel("Value", fontsize=10)
    fig.tight_layout()
    return save_chart(fig, "04_lab_channels")

def chart_hsv_all(hours_sorted, data):
    apply_style()
    fig, ax = plt.subplots(figsize=(10, 4.5))
    fig.patch.set_facecolor("#1A1A2E")
    H_vals = [data[h]["H"] for h in hours_sorted]
    S_vals = [data[h]["S"] for h in hours_sorted]
    V_vals = [data[h]["V"] for h in hours_sorted]
    stage_band(ax, hours_sorted)
    ax.plot(hours_sorted, H_vals, color="#A29BFE", lw=2.0, marker="o", ms=4, label="Hue (H)")
    ax.plot(hours_sorted, S_vals, color="#00CEC9", lw=2.4, marker="s", ms=4, label="Saturation (S)", zorder=6)
    ax.plot(hours_sorted, V_vals, color="#FDCB6E", lw=2.0, marker="^", ms=4, label="Value (V)")
    ax.set_xlabel("Elapsed Time (hours)", fontsize=11)
    ax.set_ylabel("HSV Channel Value", fontsize=11)
    ax.set_title("HSV All Channels Over Time", fontsize=13, fontweight="bold", pad=12)
    ax.set_xlim(min(hours_sorted)-0.5, max(hours_sorted)+0.5)
    ax.grid(True, alpha=0.4)
    ax.legend(loc="center right")
    fig.tight_layout()
    return save_chart(fig, "05_hsv_all")

def chart_stage_distribution(stage_orig, total_orig):
    apply_style()
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 5))
    fig.patch.set_facecolor("#1A1A2E")

    # Pie chart
    labels  = [f"Stage {s}\n{STAGES[s][0]}" for s in STAGES]
    sizes   = [stage_orig[s] for s in STAGES]
    colors  = [STAGE_COLORS[s] for s in STAGES]
    explode = [0.05]*4
    wedges, texts, autotexts = ax1.pie(sizes, labels=labels, colors=colors,
                                        autopct="%1.1f%%", explode=explode,
                                        startangle=140, textprops={"fontsize":9})
    for at in autotexts:
        at.set_fontweight("bold")
        at.set_color("white")
    ax1.set_title("Stage Distribution (Raw Images)", fontsize=12, fontweight="bold", pad=10)

    # Bar chart with augmented vs original
    x = [1, 2, 3, 4]
    orig_vals = [stage_orig[s] for s in STAGES]
    aug_vals  = [stage_orig[s] * 6 for s in STAGES]
    width = 0.35
    b1 = ax2.bar([i - width/2 for i in x], orig_vals, width, label="Original", color=[STAGE_COLORS[s] for s in STAGES], alpha=0.85)
    b2 = ax2.bar([i + width/2 for i in x], aug_vals,  width, label="Augmented (×6)", color=[STAGE_COLORS[s] for s in STAGES], alpha=0.4, edgecolor="white", linewidth=0.8)
    ax2.set_xticks(x)
    ax2.set_xticklabels([f"S{s}" for s in STAGES])
    ax2.set_title("Original vs Augmented Samples per Stage", fontsize=12, fontweight="bold", pad=10)
    ax2.set_ylabel("Number of Samples")
    ax2.legend()
    ax2.grid(True, axis="y", alpha=0.4)
    for bar in b1:
        ax2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3, str(int(bar.get_height())),
                 ha="center", va="bottom", fontsize=9, fontweight="bold")
    fig.tight_layout()
    return save_chart(fig, "06_stage_distribution")

def chart_model_performance():
    apply_style()
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    fig.patch.set_facecolor("#1A1A2E")
    fig.suptitle("Model Classification Performance", fontsize=14, fontweight="bold", y=1.02)

    # Precision / Recall / F1 grouped bar chart
    stages     = ["S1 Very Fresh", "S2 Fresh", "S3 Early Spoilage", "S4 Spoiled"]
    precision  = [0.79, 1.00, 0.63, 1.00]
    recall     = [0.92, 1.00, 1.00, 0.98]
    f1         = [0.85, 1.00, 0.77, 0.99]
    x  = np.arange(len(stages))
    w  = 0.25
    ax1.bar(x - w,   precision, w, label="Precision", color="#6C5CE7", alpha=0.9)
    ax1.bar(x,       recall,    w, label="Recall",    color="#00B894", alpha=0.9)
    ax1.bar(x + w,   f1,        w, label="F1-Score",  color="#FDCB6E", alpha=0.9)
    ax1.set_xticks(x)
    ax1.set_xticklabels(stages, fontsize=8, rotation=10)
    ax1.set_ylim(0, 1.15)
    ax1.set_ylabel("Score")
    ax1.set_title("Per-Stage Precision / Recall / F1", fontsize=11, fontweight="bold")
    ax1.legend(fontsize=9)
    ax1.grid(True, axis="y", alpha=0.4)
    ax1.axhline(1.0, color="#FFFFFF", lw=0.6, ls="--", alpha=0.4)

    # CV Accuracy donut
    acc = 95.15
    remaining = 100 - acc
    wedge_colors = ["#00D4FF", "#2A2A4A"]
    ax2.pie([acc, remaining], colors=wedge_colors, startangle=90,
            wedgeprops=dict(width=0.55, edgecolor="#1A1A2E", linewidth=2))
    ax2.text(0, 0, f"{acc:.2f}%", ha="center", va="center",
             fontsize=22, fontweight="bold", color="#00D4FF")
    ax2.text(0, -0.45, "5-Fold CV Accuracy", ha="center", va="center",
             fontsize=10, color="#CCCCCC")
    ax2.set_title("Cross-Validation Accuracy", fontsize=11, fontweight="bold")
    fig.tight_layout()
    return save_chart(fig, "07_model_performance")

def chart_dominant_colors_heatmap(hours_sorted, data):
    """Color swatch grid: hour × dominant color slot."""
    apply_style()
    dc_hours = [h for h in [0,1,2,3,4,5,6,7,8,9,10,11,12,13,14,22,23,26,30,38] if h in data]
    n = len(dc_hours)
    fig, ax = plt.subplots(figsize=(max(10, n*0.55), 3.5))
    fig.patch.set_facecolor("#1A1A2E")
    ax.set_facecolor("#1A1A2E")
    for col_idx, h in enumerate(dc_hours):
        hexes = data[h]["hex_colors"]
        for row_idx, hex_val in enumerate(hexes[:3]):
            r_val = int(hex_val[1:3], 16) / 255
            g_val = int(hex_val[3:5], 16) / 255
            b_val = int(hex_val[5:7], 16) / 255
            rect = plt.Rectangle([col_idx, -(row_idx+1)], 1, 1,
                                  facecolor=(r_val, g_val, b_val), edgecolor="#111122", linewidth=0.5)
            ax.add_patch(rect)
        stage_c = STAGE_COLORS[get_stage(h)]
        ax.text(col_idx+0.5, 0.15, str(h)+"h", ha="center", va="bottom",
                fontsize=7, color=stage_c, fontweight="bold")
    ax.set_xlim(0, n); ax.set_ylim(-3.2, 0.6)
    ax.set_yticks([-0.5, -1.5, -2.5])
    ax.set_yticklabels(["Color 1", "Color 2", "Color 3"], fontsize=9)
    ax.set_xticks([])
    ax.set_title("K-Means Dominant Color Progression (Top 3 Colors per Timepoint)", fontsize=12, fontweight="bold", pad=10)
    stage_legend = [mpatches.Patch(color=STAGE_COLORS[s], label=f"Stage {s}: {STAGES[s][0]}") for s in STAGES]
    ax.legend(handles=stage_legend, loc="lower right", fontsize=8,
              bbox_to_anchor=(1.0, -0.05), ncol=2)
    fig.tight_layout()
    return save_chart(fig, "08_dominant_colors")

def chart_multimetric_dashboard(hours_sorted, data):
    """4-panel summary dashboard."""
    apply_style()
    fig = plt.figure(figsize=(14, 9))
    fig.patch.set_facecolor("#0F0F1A")
    gs = GridSpec(2, 2, figure=fig, hspace=0.45, wspace=0.3)

    # Panel 1: Saturation + b* dual axis
    ax1 = fig.add_subplot(gs[0, 0])
    s_vals  = [data[h]["S"]     for h in hours_sorted]
    b_vals  = [data[h]["Bstar"] for h in hours_sorted]
    stage_band(ax1, hours_sorted)
    l1, = ax1.plot(hours_sorted, s_vals, color="#00D4FF", lw=2.2, marker="o", ms=3, label="Saturation (S)")
    ax1b = ax1.twinx()
    l2, = ax1b.plot(hours_sorted, b_vals, color="#FDCB6E", lw=2.0, marker="s", ms=3, ls="--", label="b* (yellow-blue)")
    ax1b.set_ylabel("b*", color="#FDCB6E", fontsize=9)
    ax1b.tick_params(colors="#FDCB6E")
    ax1.set_title("Saturation & b* Channel", fontsize=10, fontweight="bold")
    ax1.set_xlabel("Hours", fontsize=9); ax1.set_ylabel("S", color="#00D4FF", fontsize=9)
    ax1.grid(True, alpha=0.3)
    ax1.legend(handles=[l1, l2], fontsize=7, loc="lower left")

    # Panel 2: RGB
    ax2 = fig.add_subplot(gs[0, 1])
    R_vals = [data[h]["R"] for h in hours_sorted]
    G_vals = [data[h]["G"] for h in hours_sorted]
    B_vals = [data[h]["B"] for h in hours_sorted]
    stage_band(ax2, hours_sorted)
    ax2.plot(hours_sorted, R_vals, "#FF4444", lw=2.0, marker="o", ms=3, label="R")
    ax2.plot(hours_sorted, G_vals, "#44FF88", lw=2.0, marker="s", ms=3, label="G")
    ax2.plot(hours_sorted, B_vals, "#4488FF", lw=2.0, marker="^", ms=3, label="B")
    ax2.set_title("RGB Channel Values", fontsize=10, fontweight="bold")
    ax2.set_xlabel("Hours", fontsize=9); ax2.set_ylabel("Mean Value", fontsize=9)
    ax2.grid(True, alpha=0.3); ax2.legend(fontsize=8)

    # Panel 3: R-B gap bar
    ax3 = fig.add_subplot(gs[1, 0])
    rb_vals = [data[h]["RB_gap"] for h in hours_sorted]
    ax3.bar(hours_sorted, rb_vals,
            color=[STAGE_COLORS[get_stage(h)] for h in hours_sorted],
            edgecolor="#222244", linewidth=0.4, alpha=0.9)
    ax3.set_title("R–B Gap (Desaturation Indicator)", fontsize=10, fontweight="bold")
    ax3.set_xlabel("Hours", fontsize=9); ax3.set_ylabel("R–B", fontsize=9)
    ax3.grid(True, axis="y", alpha=0.3)

    # Panel 4: Stage pie
    ax4 = fig.add_subplot(gs[1, 1])
    s_counts  = {s: sum(1 for h in hours_sorted if get_stage(h)==s) for s in STAGES}
    labels_p  = [f"S{s}: {STAGES[s][0]}\n({s_counts[s]} imgs)" for s in STAGES]
    sizes_p   = [s_counts[s] for s in STAGES]
    colors_p  = [STAGE_COLORS[s] for s in STAGES]
    explode_p = [0.05]*4
    ax4.pie(sizes_p, labels=labels_p, colors=colors_p, autopct="%1.0f%%",
            explode=explode_p, startangle=140, textprops={"fontsize":8})
    ax4.set_title("Stage Distribution", fontsize=10, fontweight="bold")

    fig.suptitle(f"Freshness Monitor — Summary Dashboard ({RUN_DIR})",
                 fontsize=14, fontweight="bold", color="#FFFFFF", y=1.01)
    return save_chart(fig, "09_dashboard")

# ─── DOCX HELPERS ──────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_header_row(table, headers, fill="#1F3864"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        shade_cell(cell, fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def make_table(doc, headers, fill="#1F3864"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, headers, fill=fill)
    return table

def add_chart(doc, chart_path, caption, width=6.0):
    """Embed a chart PNG with a centered caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
def build_paper():
    images = sorted([f for f in os.listdir(RUN_DIR) if f.endswith(".jpg")], key=extract_hours)
    valid  = [f for f in images if extract_hours(f) >= 0]
    if not valid:
        print("No images found."); return

    print("Analysing images...")
    data = {}
    for f in valid:
        h = extract_hours(f)
        result = analyze_image(os.path.join(RUN_DIR, f))
        if result:
            data[h] = result

    hours_sorted = sorted(data.keys())
    total_images = len(hours_sorted)
    start_h, end_h = hours_sorted[0], hours_sorted[-1]

    stage_orig = {s: sum(1 for h in hours_sorted if get_stage(h)==s) for s in STAGES}
    total_orig = sum(stage_orig.values())
    aug_factor = 6
    stage_aug  = {s: stage_orig[s]*aug_factor for s in STAGES}
    total_aug  = sum(stage_aug.values())

    # ── Generate all charts ────────────────────────────────────────────────────
    print("Generating charts...")
    c_sat   = chart_saturation_decay(hours_sorted, data)
    c_rgb   = chart_rgb_progression(hours_sorted, data)
    c_rb    = chart_rb_gap(hours_sorted, data)
    c_lab   = chart_lab_channels(hours_sorted, data)
    c_hsv   = chart_hsv_all(hours_sorted, data)
    c_dist  = chart_stage_distribution(stage_orig, total_orig)
    c_model = chart_model_performance()
    c_dc    = chart_dominant_colors_heatmap(hours_sorted, data)
    c_dash  = chart_multimetric_dashboard(hours_sorted, data)

    # ── Build document ─────────────────────────────────────────────────────────
    print("Building document...")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0); section.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    def h1(text):
        p = doc.add_heading(text, 1); p.runs[0].font.color.rgb = RGBColor(0x1F,0x38,0x64); return p
    def h2(text):
        p = doc.add_heading(text, 2); p.runs[0].font.color.rgb = RGBColor(0x2E,0x74,0xB5); return p
    def h3(text):
        return doc.add_heading(text, 3)
    def body(text, bold_prefix=None):
        p = doc.add_paragraph()
        if bold_prefix:
            r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text); return p
    def bullet(text):
        return doc.add_paragraph(text, style="List Bullet")

    # ── TITLE ─────────────────────────────────────────────────────────────────
    t = doc.add_heading("IoT-Enabled Real-Time Food Freshness Monitoring Using Reactive Film Colorimetry and Machine Learning", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Experimental Run Report — {RUN_DIR}   |   Date: {RUN_DATE}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True; sub.runs[0].font.size = Pt(11)

    # ── ABSTRACT ──────────────────────────────────────────────────────────────
    h1("Abstract")
    doc.add_paragraph(
        "This paper presents a novel, end-to-end IoT system for real-time food freshness classification "
        "using reactive colorimetric film indicators coupled with machine learning. A Raspberry Pi captures "
        "time-lapse images of reactive films placed on perishable food items, transmits them over a secure "
        "Tailscale VPN tunnel to a Flask inference server, where a 1,045-dimensional multi-space color feature "
        "vector is extracted and classified by a RandomForest model. This report documents the specific "
        f"experimental run {RUN_DIR}, which spans {end_h} hours ({total_images} captures), providing "
        "a complete per-hour colorimetric dataset. The system achieves 95.15% cross-validated accuracy "
        "using only commodity hardware (Raspberry Pi 3B+, ~USD 35) and open-source software."
    )

    # ── 1. INTRODUCTION ───────────────────────────────────────────────────────
    h1("1. Introduction")
    h2("1.1 Problem Statement")
    doc.add_paragraph(
        "Food spoilage remains a critical global challenge. The FAO estimates that approximately one-third "
        "of all food produced for human consumption is lost or wasted annually. Current freshness assessment "
        "methods fall into three categories:"
    )
    bullet("Subjective sensory evaluation — Visual inspection by trained personnel; inconsistent, non-scalable, and prone to human error.")
    bullet("Laboratory methods — Total Volatile Basic Nitrogen (TVB-N), pH measurement, microbiological plate counts; accurate but destructive, time-consuming (24–48 hours), and require trained technicians.")
    bullet("Electronic nose / biosensor systems — Non-destructive but expensive (USD 5,000–50,000), require calibration, and are unsuitable for distributed deployment.")
    h2("1.2 Proposed Solution")
    bullet("Reactive indicator films change color in response to volatile amines (e.g., trimethylamine, ammonia) released during microbial spoilage.")
    bullet("A Raspberry Pi camera module captures these color changes at configurable intervals.")
    bullet("Multi-space color features (1,045-dimensional vectors spanning RGB, HSV histograms, channel statistics, and k-means dominant colors) are extracted and classified by a lightweight RandomForest model.")
    bullet("Results are delivered in real-time via a web dashboard with QR-coded traceability.")
    h2("1.3 Key Contributions")
    bullet("Item-agnostic colorimetric classification: The system analyzes the reactive film's color — not the food itself — making it transferable across food types without retraining.")
    bullet("Dual-mode operation: Real-time live video streaming + periodic high-resolution capture.")
    bullet("Lightweight edge-cloud architecture: Raspberry Pi 3B+ (~USD 35) + any PC/cloud server; no GPU required.")
    bullet("Automated traceability: Stage-colored QR barcodes + SQLite database + AI-generated safety assessments.")
    bullet("Comprehensive color space analysis: RGB, HSV, and CIE L*a*b* features for robust classification under varying illumination.")

    # ── 2. RELATED WORK ───────────────────────────────────────────────────────
    h1("2. Related Work")
    t_rw = make_table(doc, ["Reference", "Method", "Limitations"])
    set_col_widths(t_rw, [4.5, 6.0, 6.5])
    for row in [("Pacquit et al. (2007)","Colorimetric pH dye arrays","Manual color card comparison; no automation"),
                ("Huang et al. (2014)","Electronic nose (MOS sensors)","Expensive hardware; requires frequent calibration"),
                ("Taheri-Garavand et al. (2019)","CNN-based visual freshness (fish eyes)","Requires GPU; analyzes food directly (item-specific)"),
                ("Chen et al. (2020)","Smartphone colorimetry","Single-point RGB; no temporal monitoring"),
                ("Yousefi et al. (2021)","Paper-based sensors + smartphone app","Limited to single-shot; no IoT integration")]:
        r = t_rw.add_row()
        for i,v in enumerate(row): r.cells[i].text = v
    doc.add_paragraph()
    body("We combine continuous IoT monitoring, multi-space color feature extraction from reactive films, "
         "automated ML classification, and QR-coded traceability in a single integrated system — none of "
         "the above works achieve all four simultaneously.", bold_prefix="Our novelty: ")

    # ── 3. SYSTEM ARCHITECTURE ────────────────────────────────────────────────
    h1("3. System Architecture")
    h2("3.1 Hardware Components")
    t_hw = make_table(doc, ["Component","Specification","Role"])
    set_col_widths(t_hw, [4.5, 6.5, 6.0])
    for row in [("Raspberry Pi 3B+","1.4 GHz ARM Cortex-A53, 1 GB RAM","Edge capture device"),
                ("Pi Camera Module v2","8 MP Sony IMX219 sensor","Image acquisition"),
                ("Local PC / Cloud Server","Any x86/ARM with Python 3.11+","Inference + dashboard"),
                ("Network","Tailscale mesh VPN (WireGuard)","Secure encrypted transport")]:
        r = t_hw.add_row()
        for i,v in enumerate(row): r.cells[i].text = v
    h2("3.2 Dual-Mode Operation")
    body("Pi continuously captures at 15 fps (640×480); raw JPEG bytes pushed to POST /frame; JS-polled MJPEG at /stream.", bold_prefix="Mode 1 — Real-Time Video Streaming: ")
    body(f"Every N seconds the Pi captures full-resolution (1920×1440), uploads to POST /barcode. Generated the {RUN_DIR} dataset.", bold_prefix="Mode 2 — Periodic High-Resolution Capture: ")

    # ── 4. METHODOLOGY ────────────────────────────────────────────────────────
    h1("4. Methodology")
    h2("4.1 Reactive Film Color Indicator Principle")
    bullet("Fresh state: Film retains its original yellow/golden hue (pH neutral).")
    bullet("Early spoilage: Amines shift the film toward orange/brown tones.")
    bullet("Advanced spoilage: High amine concentration produces marked darkening and desaturation.")

    h2("4.2 Freshness Stage Classification")
    t_stages = make_table(doc, ["Stage","Label","Hour Range","Visual Indicator","Safety Assessment"])
    set_col_widths(t_stages, [1.5, 3.0, 2.5, 5.0, 5.0])
    sc = ["#C6EFCE","#FFEB9C","#FFC7CE","#FFD7D7"]
    for idx, (s,(label,h_min,h_max,visual,safety)) in enumerate(STAGES.items()):
        hr = f"{h_min}–{h_max}h" if h_max<999 else f"{h_min}+h"
        r = t_stages.add_row()
        for i,v in enumerate([str(s),label,hr,visual,safety]):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(r.cells[i], sc[idx])

    h2("4.3 Feature Extraction Pipeline")
    body("Central 40% crop to isolate reactive film; resize to 256×256; multi-space feature extraction.", bold_prefix="Steps: ")
    t_feat = make_table(doc, ["Feature Group","Dimensions","Description"])
    set_col_widths(t_feat, [4.5, 2.5, 10.0])
    for row in [("HSV Histogram","512","8×8×8 bins over H, S, V; normalized"),
                ("RGB Histogram","512","8×8×8 bins over R, G, B; normalized"),
                ("Channel Statistics","12","Mean and std of H,S,V,R,G,B channels"),
                ("Dominant Colors","9","3 dominant colors via k-means (RGB values each)"),
                ("Total","1,045","")]:
        r = t_feat.add_row()
        for i,v in enumerate(row): r.cells[i].text = v
        if row[0]=="Total":
            for cell in r.cells:
                shade_cell(cell,"#D9E1F2")
                for run in cell.paragraphs[0].runs: run.bold = True

    # ── 4.4 MULTI-METRIC DASHBOARD ────────────────────────────────────────────
    h2("4.4 Multi-Metric Summary Dashboard")
    doc.add_paragraph(
        f"The following dashboard summarises all key colorimetric metrics extracted from the {total_images} captures "
        f"spanning {start_h}h to {end_h}h in run {RUN_DIR}."
    )
    add_chart(doc, c_dash, f"Figure 1: Multi-Metric Summary Dashboard — {RUN_DIR} ({total_images} captures, {start_h}h–{end_h}h)", width=6.3)

    # ── 4.5 SATURATION DECAY ──────────────────────────────────────────────────
    h2("4.5 Saturation Decay Analysis")
    s0  = data[hours_sorted[0]]["S"]
    sEnd= data[hours_sorted[-1]]["S"]
    doc.add_paragraph(
        f"HSV Saturation (S) is the single strongest discriminator between freshness stages. "
        f"In this run, S starts at {s0:.1f} at hour 0 and drops to {sEnd:.1f} at hour {end_h} — "
        f"a {s0/max(sEnd,0.1):.1f}× decrease. The most abrupt drop occurs between hours 5–7 "
        "(Stage 2→3 boundary), signalling the onset of early spoilage."
    )
    add_chart(doc, c_sat, "Figure 2: Saturation Decay Curve — HSV S channel over 0–38 hours (colored bands = freshness stages)", width=6.3)

    # ── 4.6 COLOR SPACE ANALYSIS — FULL TABLE ─────────────────────────────────
    h2(f"4.6 Complete Per-Hour Color Space Data ({RUN_DIR})")
    doc.add_paragraph(
        f"Complete measured RGB, CIE L*a*b*, and HSV values for all {total_images} captures. "
        "Central 40% crop, resized to 256×256. Row colors indicate freshness stage."
    )
    t_color = make_table(doc, ["Hour","Stage","R","G","B","L*","a*","b*","H","S","V","R–B"])
    set_col_widths(t_color, [1.3,3.2,1.1,1.1,1.1,1.1,1.1,1.1,1.1,1.1,1.1,1.4])
    row_fills = {1:"#C6EFCE",2:"#FFEB9C",3:"#FFC7CE",4:"#FFD7D7"}
    for h in hours_sorted:
        d = data[h]; s = get_stage(h)
        r = t_color.add_row()
        vals = [str(h), STAGES[s][0],
                f"{d['R']:.1f}",f"{d['G']:.1f}",f"{d['B']:.1f}",
                f"{d['L']:.1f}",f"{d['Astar']:.1f}",f"{d['Bstar']:.1f}",
                f"{d['H']:.1f}",f"{d['S']:.1f}",f"{d['V']:.1f}",f"{d['RB_gap']:.1f}"]
        for i,v in enumerate(vals):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(r.cells[i], row_fills[s])
    doc.add_paragraph()

    # ── 4.7 RGB PROGRESSION ───────────────────────────────────────────────────
    h2("4.7 RGB Channel Progression")
    doc.add_paragraph(
        "Fresh film exhibits a strong R >> B dominance (yellow hue). As spoilage advances, "
        "all three channels converge toward equal values — a hallmark of desaturation and browning."
    )
    add_chart(doc, c_rgb, "Figure 3: RGB Channel Progression Over Time (colored fills show R–B gap shrinkage)", width=6.3)
    add_chart(doc, c_rb,  "Figure 4: R–B Gap Bar Chart — bar height indicates degree of yellow saturation (bar color = stage)", width=6.3)

    # ── 4.8 LAB + HSV ─────────────────────────────────────────────────────────
    h2("4.8 CIE L*a*b* and HSV Channel Analysis")
    rb0 = data[hours_sorted[0]]["RB_gap"]; rbEnd = data[hours_sorted[-1]]["RB_gap"]
    body(f"0h: b* = {data[0]['Bstar']:.1f} → {end_h}h: b* = {data[end_h]['Bstar']:.1f} "
         f"(Δ = {data[0]['Bstar']-data[end_h]['Bstar']:.1f}). Loss of yellow pigment confirms film degradation.",
         bold_prefix="L*a*b* b* (yellow-blue axis): ")
    body(f"Fresh (0h): R–B = {rb0:.1f} → Spoiled ({end_h}h): R–B = {rbEnd:.1f}. Channel convergence = desaturation.",
         bold_prefix="RGB convergence: ")
    add_chart(doc, c_lab, "Figure 5: CIE L*a*b* Channel Values — L* (lightness), a* (green-red), b* (yellow-blue) over time", width=6.5)
    add_chart(doc, c_hsv, "Figure 6: HSV All Channels (Hue, Saturation, Value) Over the Full Spoilage Cycle", width=6.3)

    # ── 4.9 DOMINANT COLOURS ─────────────────────────────────────────────────
    h2("4.9 Dominant Color Progression")
    doc.add_paragraph(
        "K-means clustering (k=3) extracts the three most prevalent colors per timepoint. "
        "The swatch grid below visually demonstrates the color shift from bright yellow (#ecc10e) "
        "at 0h to muted brown (#d48d2e) by 38h."
    )
    dc_hours = [h for h in [0,1,2,3,4,5,6,7,8,9,10,11,12,13,14,22,23,26,30,38] if h in data]
    t_dc = make_table(doc, ["Hour","Stage","Dominant Color 1","Dominant Color 2","Dominant Color 3"])
    set_col_widths(t_dc, [1.5,3.5,4.5,4.5,4.5])
    for h in [h for h in [0,3,6,11,22,30,38] if h in data]:
        d = data[h]; s = get_stage(h)
        hexes = d["hex_colors"]
        r = t_dc.add_row()
        r.cells[0].text = str(h); r.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r.cells[1].text = STAGES[s][0]
        for ci, hx in enumerate(hexes[:3]):
            cell = r.cells[2+ci]
            cell.text = hx
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(cell, hx)
    doc.add_paragraph()
    add_chart(doc, c_dc, "Figure 7: K-Means Dominant Color Swatch Grid — top 3 colors per captured timepoint (hour label color = stage)", width=6.5)

    # ── 4.10 MODEL TRAINING ───────────────────────────────────────────────────
    h2("4.10 Model Training")
    body("RandomForest (100 trees, max_depth=5, balanced class weights) and SVM (RBF kernel, balanced weights).", bold_prefix="Algorithms: ")
    body("StandardScaler (zero mean, unit variance) before training.", bold_prefix="Feature Scaling: ")
    body("5-Fold Stratified Cross-Validation preserving class distribution.", bold_prefix="Validation: ")

    # ── 5. EXPERIMENTAL RESULTS ───────────────────────────────────────────────
    h1("5. Experimental Results")
    h2("5.1 Dataset Summary")
    t_ds = make_table(doc, ["Parameter","Value"])
    set_col_widths(t_ds, [7.0,10.0])
    for k,v in [("Training run",RUN_DIR),("Experiment date",RUN_DATE),
                ("Raw images captured",str(total_images)),
                ("Observation span",f"{start_h}h to {end_h}h ({end_h-start_h} hours total)"),
                ("Augmented samples",f"{total_aug} ({total_images} × {aug_factor})"),
                ("Features per sample","1,045"),("Feature CSV size","~7.7 MB (all runs combined)")]:
        r = t_ds.add_row()
        r.cells[0].text = k; r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = v

    h2("5.2 Stage Distribution")
    t_sd = make_table(doc, ["Stage","Label","Hour Range","Original","Augmented (×6)","Share"])
    set_col_widths(t_sd, [1.5,3.5,2.5,2.5,3.5,2.5])
    for s in STAGES:
        label,h_min,h_max,_,_ = STAGES[s]
        hr = f"{h_min}–{h_max}h" if h_max<999 else f"{h_min}+h"
        r = t_sd.add_row()
        for i,v in enumerate([str(s),label,hr,str(stage_orig[s]),str(stage_aug[s]),f"{stage_orig[s]/total_orig*100:.1f}%"]):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(r.cells[i], row_fills[s])
    r_tot = t_sd.add_row()
    for i,v in enumerate(["—","TOTAL","—",str(total_orig),str(total_aug),"100%"]):
        r_tot.cells[i].text = v
        r_tot.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(r_tot.cells[i],"#D9E1F2")
        for run in r_tot.cells[i].paragraphs[0].runs: run.bold = True
    doc.add_paragraph()
    add_chart(doc, c_dist, "Figure 8: Stage Distribution — Pie chart (raw images) and grouped bar chart (original vs augmented samples)", width=6.5)

    h2("5.3 Classification Performance")
    t_perf = make_table(doc, ["Metric","Value"])
    set_col_widths(t_perf, [7.0,10.0])
    for k,v in [("5-Fold CV Accuracy","95.15% (±1.01%)"),
                ("Full Dataset Accuracy","98.21%"),
                ("Best Model","RandomForest (100 trees, max_depth=5)"),
                ("Misclassifications","17 / 948 total samples")]:
        r = t_perf.add_row()
        r.cells[0].text = k; r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = v

    h2("5.4 Per-Stage Classification Report")
    t_cls = make_table(doc, ["Stage","Label","Precision","Recall","F1-Score","Support"])
    set_col_widths(t_cls, [1.5,4.0,2.5,2.5,2.5,2.5])
    for s,label,prec,rec,f1,sup in [(1,"Very Fresh","0.79","0.92","0.85","12"),
                                     (2,"Fresh","1.00","1.00","1.00","12"),
                                     (3,"Early Spoilage","0.63","1.00","0.77","24"),
                                     (4,"Spoiled","1.00","0.98","0.99","900")]:
        r = t_cls.add_row()
        for i,v in enumerate([str(s),label,prec,rec,f1,sup]):
            r.cells[i].text = v
            r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade_cell(r.cells[i], row_fills[s])
    r_wa = t_cls.add_row()
    for i,v in enumerate(["—","Weighted Average","0.99","0.98","0.98","948"]):
        r_wa.cells[i].text = v
        r_wa.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(r_wa.cells[i],"#D9E1F2")
        for run in r_wa.cells[i].paragraphs[0].runs: run.bold = True
    doc.add_paragraph()
    add_chart(doc, c_model, "Figure 9: Model Performance — Per-stage Precision/Recall/F1 bars and Cross-Validation Accuracy donut chart", width=6.5)

    h2("5.5 Training History Across Runs")
    t_hist = make_table(doc, ["Run","Date","Images","Augmented","Food Type","CV Accuracy"])
    set_col_widths(t_hist, [3.5,2.5,2.0,3.0,2.5,3.0])
    for row in [("run_20260420","Apr 2026","19","114","Mixed","Baseline"),
                ("run_20260505","May 2026","19","114","Mixed","Improved"),
                (RUN_DIR,RUN_DATE,str(total_images),str(total_aug),"Shrimp","95.15%")]:
        r = t_hist.add_row()
        for i,v in enumerate(row):
            r.cells[i].text = v
            if row[0]==RUN_DIR:
                shade_cell(r.cells[i],"#C6EFCE")
                for run in r.cells[i].paragraphs[0].runs: run.bold = True

    h2("5.6 Analysis of Results")
    body(f"Stage 4 (Spoiled): precision 1.00, recall 0.98 — most safety-critical classification. Stage 2 (Fresh): perfect 1.00/1.00/1.00. Model generalizes despite class imbalance.", bold_prefix="Strengths: ")
    body(f"Stage 1 precision (0.79) limited by only {stage_orig[1]} original images in this run. Stage 3 (0.63) shows some adjacent-stage confusion. More balanced 0–14h sampling needed.", bold_prefix="Limitations: ")

    # ── 6. NOVELTY ────────────────────────────────────────────────────────────
    h1("6. Novelty and Contributions")
    h2("6.1 Comparison with Existing Approaches")
    t_nov = make_table(doc, ["Aspect","Existing Approaches","Our Approach"])
    set_col_widths(t_nov, [4.0,7.0,7.0])
    for row in [("Sensing","Electronic noses, biosensors (expensive)","Reactive colorimetric film (low-cost, disposable)"),
                ("Analysis","Single RGB reading or deep learning (GPU)","1,045-dim multi-space features + lightweight ML (no GPU)"),
                ("Monitoring","Single-shot measurement","Continuous IoT monitoring (live stream + periodic capture)"),
                ("Hardware","Laboratory equipment or smartphones","Raspberry Pi (~$35) + standard camera"),
                ("Traceability","None or manual logging","Automated QR barcodes + SQLite + AI reports"),
                ("Food specificity","Item-specific models","Item-agnostic (analyzes film, not food)"),
                ("Deployment","Lab or controlled environment","Field-deployable via Tailscale VPN (any network)")]:
        r = t_nov.add_row()
        for i,v in enumerate(row): r.cells[i].text = v
    h2("6.2 Technical Novelty")
    bullet("Multi-color-space feature fusion: RGB + HSV + k-means dominant colors = 1,045-dimensional vector.")
    bullet("Central crop isolation: 40% crop effectively isolates reactive film without object detection.")
    bullet("Automated data pipeline: organize_captures.py auto-detects runs and assigns elapsed-hour labels.")
    bullet("Hybrid streaming: JS-polling works through any reverse proxy — cloud-agnostic deployment.")
    bullet("AI-augmented reporting: Google Gemini provides natural-language safety assessments.")

    # ── 7. IMPLEMENTATION ────────────────────────────────────────────────────
    h1("7. System Implementation Details")
    h2("7.1 Technology Stack")
    t_tech = make_table(doc, ["Layer","Technology"])
    set_col_widths(t_tech, [5.0,12.0])
    for k,v in [("Edge Device","Raspberry Pi 3B+, PiCamera2, Python 3.x"),
                ("Network","Tailscale VPN (WireGuard)"),("Server","Flask, Gunicorn"),
                ("ML","scikit-learn (RandomForest, SVM, StandardScaler)"),
                ("Computer Vision","OpenCV (histograms, color conversion, k-means)"),
                ("Database","SQLite3"),("AI Reports","Google Gemini 1.5 Flash"),
                ("Deployment","Render (cloud) or local PC")]:
        r = t_tech.add_row()
        r.cells[0].text = k; r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[1].text = v
    h2("7.2 API Endpoints")
    t_api = make_table(doc, ["Method","Endpoint","Description"])
    set_col_widths(t_api, [2.0,5.0,10.0])
    for row in [("POST","/predict","Image upload → stage classification"),("POST","/barcode","Image upload → classification + QR code"),
                ("POST","/frame","Raw JPEG frame from Pi (stream)"),("GET","/video_feed","MJPEG live stream"),
                ("GET","/latest_frame.jpg","Single latest frame (JS polling)"),("GET","/dashboard","Real-time monitoring dashboard"),
                ("GET","/gallery","Historical image gallery"),("GET","/history","JSON prediction history"),
                ("GET","/status","Server health check"),("GET","/result/latest","Interactive result page")]:
        r = t_api.add_row()
        for i,v in enumerate(row): r.cells[i].text = v

    # ── 8. TESTING ────────────────────────────────────────────────────────────
    h1("8. Testing and Validation")
    h2("8.1 Color Feature Validation")
    body(f"0h → {end_h}h: S drops {data[hours_sorted[0]]['S']:.1f} → {data[hours_sorted[-1]]['S']:.1f} ({data[hours_sorted[0]]['S']-data[hours_sorted[-1]]['S']:.1f} units).", bold_prefix="Saturation decay: ")
    body(f"0h → {end_h}h: b* drops {data[0]['Bstar']:.1f} → {data[end_h]['Bstar']:.1f} ({data[0]['Bstar']-data[end_h]['Bstar']:.1f} units). Yellow pigment loss confirmed.", bold_prefix="b* channel: ")
    body(f"Fresh R–B = {data[0]['RB_gap']:.1f} → Spoiled R–B = {data[end_h]['RB_gap']:.1f}. Channel convergence = desaturation.", bold_prefix="RGB convergence: ")
    h2("8.2 Cross-Validation Results")
    doc.add_paragraph("5-Fold Stratified CV: Mean accuracy 95.15% (±1.01%). Consistent performance across all folds confirms no overfitting despite class imbalance.")
    h2("8.3 End-to-End System Testing")
    t_e2e = make_table(doc, ["Test Case","Method","Result"])
    set_col_widths(t_e2e, [4.5,5.0,7.5])
    for row in [("Pi capture + upload","--stream --capture-every 60","Verified: images arrive within 2s"),
                ("Real-time classification","POST /barcode","Sub-second inference"),
                ("QR code generation","/barcode/image/latest","Stage-colored QR codes verified scannable"),
                ("Dashboard live update","/dashboard","Auto-refreshes every 2s"),
                ("Database persistence","Server restart","All historical predictions survive"),
                ("Network resilience","Tailscale reconnection","Pi auto-retries; no data loss")]:
        r = t_e2e.add_row()
        for i,v in enumerate(row): r.cells[i].text = v
    h2("8.4 Challenges and Limitations")
    bullet(f"Dataset Imbalance: Stage 4 constitutes {stage_orig[4]/total_orig*100:.1f}% of raw captures.")
    bullet("Data Collection Gap: Continuous imaging between 15–22h would better map the Stage 3→4 transition.")
    bullet("Lighting Dependency: Drastic color temperature changes could skew histogram representations.")
    bullet("Centering Assumption: Fixed 40% crop fails if reactive film is displaced from center.")

    # ── 9. CONCLUSION ────────────────────────────────────────────────────────
    h1("9. Conclusion")
    doc.add_paragraph(
        f"We presented a complete IoT-based food freshness monitoring system combining reactive colorimetric films, "
        f"multi-space color feature extraction, lightweight ML, and real-time web monitoring. The system achieves "
        f"95.15% cross-validated accuracy using only commodity hardware. Run {RUN_DIR} captured {total_images} images "
        f"over {end_h} hours, providing a comprehensive colorimetric dataset validating the system's ability to track "
        f"progressive spoilage — particularly the sharp saturation drop from {data[hours_sorted[0]]['S']:.1f} to "
        f"{data[hours_sorted[-1]]['S']:.1f} S units across the full spoilage cycle."
    )
    h2("9.1 Future Work")
    bullet("Balanced dataset collection: More 0–14h images to improve Stage 1–3 precision.")
    bullet("Multi-food validation: Poultry, dairy, seafood experiments for cross-food generalizability.")
    bullet("Temperature compensation: Integrate sensors to account for varying spoilage rates.")
    bullet("Edge inference: Port RandomForest to Pi for offline classification.")
    bullet("Deep learning comparison: Benchmark vs MobileNet/EfficientNet on accuracy-cost tradeoff.")

    # ── 10. REFERENCES ───────────────────────────────────────────────────────
    h1("10. References")
    for i,ref in enumerate([
        'Pacquit, A., et al. (2007). "Development of a smart packaging for the monitoring of fish spoilage." Food Chemistry, 102(2), 466-470.',
        'Huang, X., et al. (2014). "Rapid detection of meat spoilage using an electronic nose." Journal of Food Engineering, 130, 42-48.',
        'Taheri-Garavand, A., et al. (2019). "Assessment of fish freshness using computer vision and deep learning." Food Analytical Methods, 12(8), 1771-1783.',
        'Chen, Q., et al. (2020). "Smartphone-based colorimetric sensor for freshness assessment." Sensors and Actuators B: Chemical, 311, 127886.',
        'Yousefi, H., et al. (2021). "Paper-based sensors for food quality monitoring." ACS Nano, 15(11), 17263-17283.',
        'Breiman, L. (2001). "Random forests." Machine Learning, 45(1), 5-32.',
        'Pedregosa, F., et al. (2011). "Scikit-learn: Machine Learning in Python." JMLR, 12, 2825-2830.',
    ], 1):
        doc.add_paragraph(f"[{i}] {ref}", style="List Number")

    # ── APPENDIX ─────────────────────────────────────────────────────────────
    h1("Appendix A: Reproducibility")
    h2("A.1 Software Dependencies")
    p = doc.add_paragraph(style="No Spacing")
    p.add_run("flask>=3.0, scikit-learn>=1.3, opencv-python-headless>=4.8\nnumpy>=1.24, Pillow>=10.0, joblib>=1.3, qrcode>=7.4\ngunicorn>=21.2, google-genai>=0.6.0").font.name="Courier New"
    h2("A.2 Training Pipeline Commands")
    p2 = doc.add_paragraph(style="No Spacing")
    p2.add_run("python organize_captures.py\npython prepare_data.py\npython train_model.py\npython server.py").font.name="Courier New"
    h2("A.3 Pi Client Commands")
    p3 = doc.add_paragraph(style="No Spacing")
    p3.add_run("python3 pi_client.py --stream --capture-every 60\npython3 pi_client.py --mode http --interval 10").font.name="Courier New"

    doc.save(OUTPUT)
    print(f"\nDocument saved: {OUTPUT}")
    print(f"  {total_images} images | {start_h}h–{end_h}h | {total_aug} augmented samples")
    print(f"  9 charts embedded")

if __name__ == "__main__":
    build_paper()

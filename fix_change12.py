"""Fix Change 12 in the revised manuscript."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

OUTPUT_PATH = r'C:\Users\ACER\Downloads\Manuscript Food chemistry - REVISED.docx'
doc = Document(OUTPUT_PATH)

# Find and fix the exact phrase
for i, para in enumerate(doc.paragraphs):
    if "intelligent (pH-responsive colorimetric sensing) packaging functions" in para.text:
        text = para.text
        new_text = text.replace(
            "intelligent (pH-responsive colorimetric sensing) packaging functions",
            "intelligent packaging functions, including pH-responsive colorimetric sensing"
        )
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        print(f"Fixed PARA {i}: replacement applied")
        break
else:
    print("ERROR: Target phrase not found!")

doc.save(OUTPUT_PATH)
print("Document saved.")

# Verify
doc2 = Document(OUTPUT_PATH)
for para in doc2.paragraphs:
    if "including pH-responsive colorimetric sensing" in para.text:
        print("VERIFICATION: PASS - Change 12 now applied correctly")
        break
else:
    print("VERIFICATION: FAIL")

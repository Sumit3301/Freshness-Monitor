"""Debug and fix Change 12 in the revised manuscript."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

OUTPUT_PATH = r'C:\Users\ACER\Downloads\Manuscript Food chemistry - REVISED.docx'
doc = Document(OUTPUT_PATH)

# Find the conclusions paragraph
for i, para in enumerate(doc.paragraphs):
    if "packaging functions" in para.text and "intelligent" in para.text:
        print(f"PARA {i}: ...{para.text[para.text.index('intelligent')-20:para.text.index('packaging functions')+30]}...")
        
# Also search for the exact phrase
for i, para in enumerate(doc.paragraphs):
    if "pH-responsive colorimetric sensing" in para.text:
        # Find the context around it
        idx = para.text.index("pH-responsive colorimetric sensing")
        start = max(0, idx-50)
        end = min(len(para.text), idx+80)
        print(f"\nFOUND in PARA {i}: ...{para.text[start:end]}...")

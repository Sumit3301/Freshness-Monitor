import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob

def create_research_paper(run_dir, output_file):
    document = Document()

    # Title
    title = document.add_heading('Freshness Monitor: Spoilage Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    document.add_heading('1. Introduction', level=1)
    p = document.add_paragraph(
        'This document presents the visual findings and data analysis for the '
        f'experiment run stored in the directory: {run_dir}. The primary objective '
        'is to monitor and document the progressive degradation and spoilage of '
        'the observed item over a specific time duration.'
    )

    # Dataset Summary
    document.add_heading('2. Dataset Summary', level=1)
    
    images = glob.glob(os.path.join(run_dir, '*.jpg'))
    
    def extract_hours(filename):
        basename = os.path.basename(filename)
        try:
            return int(basename.replace('h.jpg', ''))
        except ValueError:
            return -1

    images.sort(key=extract_hours)
    valid_images = [img for img in images if extract_hours(img) != -1]

    if not valid_images:
        document.add_paragraph("No valid image data found in the specified directory.")
        document.save(output_file)
        return

    start_time = extract_hours(valid_images[0])
    end_time = extract_hours(valid_images[-1])
    total_images = len(valid_images)

    p2 = document.add_paragraph()
    p2.add_run(f'Total Captures: ').bold = True
    p2.add_run(f'{total_images} images\n')
    p2.add_run(f'Observation Duration: ').bold = True
    p2.add_run(f'From {start_time}h to {end_time}h\n')

    # Visual Evidence
    document.add_heading('3. Visual Spoilage Progression', level=1)
    document.add_paragraph('The following images illustrate the physical changes observed over time.')

    # Select representative images (Start, Middle, End)
    indices = [0, total_images // 2, total_images - 1]
    selected_images = []
    for idx in indices:
        if idx < total_images and valid_images[idx] not in selected_images:
            selected_images.append(valid_images[idx])

    for img_path in selected_images:
        basename = os.path.basename(img_path)
        hour = extract_hours(img_path)
        
        document.add_heading(f'Observation at {hour} Hours ({basename})', level=2)
        try:
            document.add_picture(img_path, width=Inches(4.0))
        except Exception as e:
            document.add_paragraph(f'[Error loading image {basename}: {e}]')

    # Conclusion
    document.add_heading('4. Conclusion', level=1)
    document.add_paragraph(
        'The captured progression clearly demonstrates the expected stages of spoilage. '
        'Further analysis using the machine learning pipeline (RandomForest/SVM) on the extracted '
        'features (color histograms, texture patterns) from these captures will provide a more '
        'quantitative evaluation of freshness.'
    )

    document.save(output_file)
    print(f"Research paper successfully generated and saved to {output_file}")

if __name__ == '__main__':
    run_directory = 'run_20260503_1136'
    output_filename = f'Research_Paper_{run_directory}.docx'
    create_research_paper(run_directory, output_filename)

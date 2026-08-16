import docx
from docx import Document

doc = Document(r'C:\Users\ACER\Downloads\comments review.docx')

with open(r'd:\POC project\comments_full.txt', 'w', encoding='utf-8') as f:
    f.write("=== COMMENTS REVIEW DOCUMENT ===\n")
    f.write(f"Total paragraphs: {len(doc.paragraphs)}\n\n")

    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "None"
            f.write(f"PARA {i} [Style: {style}]:\n")
            f.write(f"  {text}\n\n")

    f.write("\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\nTable {t_idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns\n")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    f.write(f"  Row {r_idx}, Col {c_idx}: {cell_text[:500]}\n")

print("Done writing comments")
